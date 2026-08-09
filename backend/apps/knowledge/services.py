import hashlib
import logging
import re
import uuid
import zipfile
from xml.etree import ElementTree
from pathlib import Path
from django.conf import settings
from django.db import transaction
from django.utils import timezone
from common.exceptions import BusinessError
from apps.ai_service.services import get_config
from .models import ParseTask, TextChunk

logger = logging.getLogger("knowledge")

MARKDOWN_IMAGE_PATTERN = re.compile(r"!\[[^\]]*\]\([^\n)]*\)")
HTML_IMAGE_PATTERN = re.compile(r"<img\b[^>]*>", re.IGNORECASE)
BASE64_RUN_PATTERN = re.compile(r"[A-Za-z0-9+/]{100,}={0,2}")

# Some Chinese textbook PDFs use GBK user-defined character slots for Latin
# variables. PDF extractors then return look-alike CJK characters such as
# "狓" instead of "x". The second GBK byte still preserves the original
# alphabet position, so the mapping is deterministic and does not affect
# normal Chinese characters such as "独" or "状".
PDF_SYMBOL_REPLACEMENTS = {
    "\ue010": ".",
    "\ue011": "-",
    "\ue012": "*",
    "\ue01b": "∈",
    "\ue020": "⊆",
    "\ue021": "⊇",
    "\ue02f": "∀",
    "\ue039": "⇔",
    "\ue03c": "⇒",
    "\ue047": "\t",
    "\ue055": "∃",
    "\ue05b": "∉",
    "\ue061": "⊄",
    "\ue07e": "∅",
    "\ue0b4": "▱",
    # This glyph is the second half of an already extracted vector arrow.
    "\ue5c6": "",
}


def normalize_pdf_text(text):
    normalized = []
    for char in str(text or ""):
        replacement = PDF_SYMBOL_REPLACEMENTS.get(char)
        if replacement is not None:
            normalized.append(replacement)
            continue
        try:
            encoded = char.encode("gbk")
        except UnicodeEncodeError:
            encoded = b""
        if len(encoded) == 2 and encoded[0] == 0xA0:
            if 0xC1 <= encoded[1] <= 0xDA:
                normalized.append(chr(ord("A") + encoded[1] - 0xC1))
                continue
            if 0xE1 <= encoded[1] <= 0xFA:
                normalized.append(chr(ord("a") + encoded[1] - 0xE1))
                continue
        # Decorative exercise icons from the embedded symbol fonts are often
        # exposed as Ethiopic code points and add noise to RAG text.
        if 0x1200 <= ord(char) <= 0x13FF:
            continue
        normalized.append(char)
    return "".join(normalized)


def is_meaningful_chunk(content):
    """过滤无法支持命题的图片路径、Base64 数据、纯目录和纯索引块。"""
    original = str(content or "")
    if not original.strip():
        return False
    base64_chars = sum(len(match.group(0)) for match in BASE64_RUN_PATTERN.finditer(original))
    if base64_chars >= max(100, len(original) * 0.2):
        return False
    cleaned = MARKDOWN_IMAGE_PATTERN.sub(" ", original)
    cleaned = HTML_IMAGE_PATTERN.sub(" ", cleaned)
    cleaned = BASE64_RUN_PATTERN.sub(" ", cleaned)
    readable = re.sub(r"[^A-Za-z0-9\u4e00-\u9fff]+", "", cleaned)
    if len(readable) < 20:
        return False
    dotted_toc_lines = len(re.findall(r"(?:\.\s*){5,}\s*\d+", cleaned))
    index_references = len(re.findall(r"(?<=[,，])\s*\d+", cleaned))
    sentence_marks = len(re.findall(r"[.!?。！？；;]", cleaned))
    if dotted_toc_lines >= 2:
        return False
    if index_references >= 8 and sentence_marks < 3:
        return False
    return True

class TaskYield(Exception):
    """向量任务主动让出Worker，不代表处理失败。"""

class BaseParser:
    def parse(self, path): raise NotImplementedError

class PdfParser(BaseParser):
    def parse(self, path):
        try:
            import fitz
            document = fitz.open(path)
            pages = [{"text": normalize_pdf_text(page.get_text("text")), "page": i + 1} for i, page in enumerate(document)]
            if not any(x["text"].strip() for x in pages):
                raise BusinessError("该PDF可能是扫描文件，当前版本暂不支持OCR识别。", 40021)
            return pages
        except BusinessError: raise
        except Exception as exc: raise BusinessError(f"PDF解析失败：{exc}", 40022) from exc

class DocxParser(BaseParser):
    def parse(self, path):
        try:
            from docx import Document
            doc = Document(path)
            blocks = [p.text for p in doc.paragraphs]
            for table in doc.tables:
                blocks.extend("\t".join(cell.text for cell in row.cells) for row in table.rows)
            return [{"text": "\n".join(blocks), "page": None}]
        except KeyError as exc:
            # 某些文档生成器会留下 Target="../NULL" 等无效图片关系。
            # 图片缺失不应阻止正文入库，因此直接读取完整的 document.xml 文本。
            try:
                with zipfile.ZipFile(path) as package:
                    root = ElementTree.fromstring(package.read("word/document.xml"))
                namespace = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
                blocks = []
                for paragraph in root.iter(f"{namespace}p"):
                    parts = []
                    for node in paragraph.iter():
                        if node.tag == f"{namespace}t" and node.text:
                            parts.append(node.text)
                        elif node.tag == f"{namespace}tab":
                            parts.append("\t")
                        elif node.tag in {f"{namespace}br", f"{namespace}cr"}:
                            parts.append("\n")
                    text = "".join(parts).strip()
                    if text:
                        blocks.append(text)
                if not blocks:
                    raise ValueError("document.xml 中没有可提取的正文")
                logger.warning("DOCX包含缺失附件，已忽略附件并提取正文：%s（%s）", path, exc)
                return [{"text": "\n".join(blocks), "page": None}]
            except Exception as fallback_exc:
                raise BusinessError(f"Word文件解析失败，文件可能已损坏：{fallback_exc}", 40023) from fallback_exc
        except Exception as exc:
            raise BusinessError(f"Word文件解析失败，文件可能已损坏：{exc}", 40023) from exc

class TextParser(BaseParser):
    def parse(self, path):
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try: return [{"text": Path(path).read_text(encoding=encoding), "page": None}]
            except UnicodeDecodeError: continue
        raise BusinessError("文本文件编码无法识别，请转换为UTF-8后重试。", 40024)

class ParserFactory:
    @staticmethod
    def create(file_type):
        mapping = {"pdf": PdfParser, "docx": DocxParser, "txt": TextParser, "md": TextParser, "markdown": TextParser}
        parser = mapping.get(file_type.lower())
        if not parser: raise BusinessError("当前文件类型不支持解析。", 40020)
        return parser()

class TextCleaner:
    @staticmethod
    def clean(text):
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
        text = text.replace("\r\n", "\n").replace("\r", "\n")
        # Markdown 资料常把图片路径甚至 Base64 原始数据写进正文。
        # 这些内容不能用于出题，还会干扰 Embedding 检索。
        text = MARKDOWN_IMAGE_PATTERN.sub(" ", text)
        text = HTML_IMAGE_PATTERN.sub(" ", text)
        text = BASE64_RUN_PATTERN.sub(" ", text)
        text = re.sub(r"(?m)^\s*(?:第?\s*\d+\s*页|[-—]?\s*\d+\s*[-—]?)\s*$", "", text)
        text = re.sub(r"(?<![。！？：；\n])\n(?=[\u4e00-\u9fff])", "", text)
        text = re.sub(r"[ \t]+", " ", text)
        lines, seen = [], set()
        for line in text.splitlines():
            value = line.strip()
            if not value: continue
            marker_free = re.sub(r"^#{1,6}\s*", "", value)
            if marker_free in seen: continue
            seen.add(marker_free); lines.append(marker_free)
        return "\n".join(lines)

def split_text(text, chunk_size=800, overlap=120):
    if chunk_size <= 0 or overlap < 0 or overlap >= chunk_size:
        raise ValueError("分块长度必须大于0，重叠长度必须小于分块长度")
    chunks, start = [], 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        if end < len(text):
            candidates = [text.rfind(mark, start + chunk_size // 2, end) for mark in ("\n", "。", "；", "！", "？")]
            boundary = max(candidates)
            if boundary > start: end = boundary + 1
        value = text[start:end].strip()
        if value: chunks.append(value)
        if end >= len(text): break
        start = max(start + 1, end - overlap)
    return chunks

class VectorService:
    # v3 is stored under an ASCII-only path. Chroma's Windows HNSW backend can
    # write into a Unicode path but may fail to persist/reopen its binary files.
    collection_name = "knowledge_chunks_v3"
    def __init__(self):
        self.config = get_config()
    def _collection(self):
        try:
            import chromadb
            from chromadb.config import Settings
            from langchain_ollama import OllamaEmbeddings
            client = chromadb.PersistentClient(
                path=str(settings.CHROMA_PATH),
                settings=Settings(anonymized_telemetry=False),
            )
            collection = client.get_or_create_collection(self.collection_name, metadata={"hnsw:space": "cosine"})
            embeddings = OllamaEmbeddings(
                model=self.config["embedding_model"],
                base_url=self.config["ollama_base_url"],
                sync_client_kwargs={"timeout": float(self.config["timeout"])},
                keep_alive=int(self.config.get("keep_alive", 900)),
            )
            return collection, embeddings
        except ImportError as exc: raise BusinessError("向量数据库依赖未安装，请执行 pip install -r requirements.txt。", 50021, 500) from exc
    def index(self, chunks, progress_callback=None, cancel_callback=None, yield_callback=None):
        collection, embeddings = self._collection()
        batch_size = max(1, min(int(self.config.get("embedding_batch_size", 32)), 128))
        total = len(chunks)
        for start in range(0, total, batch_size):
            if cancel_callback and cancel_callback():
                raise BusinessError("任务已取消", 40901, 409)
            batch = chunks[start:start + batch_size]
            texts = [x.content for x in batch]
            vectors = embeddings.embed_documents(texts)
            ids = [x.vector_id for x in batch]
            metadatas = [{"course_id": x.course_id, "file_id": x.knowledge_file_id, "chapter_id": x.chapter_id or 0, "chunk_id": x.id, "page_number": x.page_number or 0} for x in batch]
            collection.upsert(ids=ids, embeddings=vectors, documents=texts, metadatas=metadatas)
            TextChunk.objects.filter(id__in=[x.id for x in batch]).update(vector_status="SUCCESS", embedding_model=self.config["embedding_model"], vector_error="")
            if progress_callback:
                progress_callback(min(start + len(batch), total), total)
            # 每完成一批才让出，确保向量和SQLite状态已经持久化。
            if yield_callback and yield_callback():
                raise TaskYield()
    def delete_file(self, file_id):
        try:
            collection, _ = self._collection()
            collection.delete(where={"file_id": int(file_id)})
        except Exception as exc:
            logger.warning("清理file_id=%s向量失败: %s", file_id, exc)
    @staticmethod
    def _where_filter(filters):
        where = {k: int(v) for k, v in filters.items() if v not in (None, "", 0)}
        if len(where) > 1:
            return {"$and": [{k: {"$eq": v}} for k, v in where.items()]}
        if len(where) == 1:
            return {next(iter(where)): {"$eq": next(iter(where.values()))}}
        return None
    def search(self, query, filters, top_k=5, threshold=0.25):
        collection, embeddings = self._collection()
        result = collection.query(query_embeddings=[embeddings.embed_query(query)], n_results=max(top_k * 3, top_k), where=self._where_filter(filters), include=["documents", "metadatas", "distances"])
        items = []
        for idx, distance in enumerate((result.get("distances") or [[]])[0]):
            similarity = max(0.0, min(1.0, 1 - float(distance)))
            if similarity < threshold: continue
            document = TextCleaner.clean(result["documents"][0][idx])
            if not is_meaningful_chunk(document): continue
            meta = result["metadatas"][0][idx]
            items.append({"chunk_id": meta.get("chunk_id"), "content": document, "similarity": round(similarity, 4), "distance": round(float(distance), 4), **meta})
            if len(items) >= top_k: break
        return items[:top_k]

    def _lexical_fallback(self, query, filters, limit):
        """当某主题的向量候选全是旧图片数据时，从SQLite原文块中补充正文。"""
        queryset = TextChunk.objects.filter(
            content__icontains=query,
            knowledge_file__is_deleted=False,
            knowledge_file__is_enabled=True,
        ).select_related("knowledge_file")
        mapping = {"course_id": "course_id", "file_id": "knowledge_file_id", "chapter_id": "chapter_id"}
        for key, field in mapping.items():
            value = filters.get(key)
            if value not in (None, "", 0):
                queryset = queryset.filter(**{field: int(value)})
        results = []
        # 先取更多命中再过滤，避免书后索引或图片占用限额。
        for chunk in queryset.order_by("id")[:300]:
            content = TextCleaner.clean(chunk.content)
            if not is_meaningful_chunk(content):
                continue
            results.append({
                "chunk_id": chunk.id,
                "content": content,
                "similarity": 0.5,
                "distance": 0.5,
                "retrieval_query": query,
                "match_type": "KEYWORD",
                "course_id": chunk.course_id,
                "file_id": chunk.knowledge_file_id,
                "chapter_id": chunk.chapter_id or 0,
                "page_number": chunk.page_number or 0,
            })
            if len(results) >= limit:
                break
        return results

    def search_many(self, queries, filters, per_query=3, threshold=0.25, max_results=50):
        """一次 Embedding 多个主题并分别检索，再轮询合并结果。

        轮询合并可以防止“Python”等宽泛主题占满整个结果池，
        保证用户列出的 CNN、SVM、PCA 等主题都有自己的 RAG 上下文。
        """
        cleaned_queries = list(dict.fromkeys(str(query).strip() for query in queries if str(query).strip()))
        if not cleaned_queries:
            return []
        collection, embeddings = self._collection()
        vectors = embeddings.embed_documents(cleaned_queries)
        result = collection.query(
            query_embeddings=vectors,
            # 某些 Markdown 书籍的图片数据已经被旧版程序存入向量库。
            # 先多取候选，再只保留少量有效正文，无需用户重新解析全部资料。
            n_results=max(int(per_query) * 25, 50),
            where=self._where_filter(filters),
            include=["documents", "metadatas", "distances"],
        )
        buckets = []
        all_distances = result.get("distances") or []
        all_documents = result.get("documents") or []
        all_metadatas = result.get("metadatas") or []
        for query_index, query in enumerate(cleaned_queries):
            bucket = []
            distances = all_distances[query_index] if query_index < len(all_distances) else []
            documents = all_documents[query_index] if query_index < len(all_documents) else []
            metadatas = all_metadatas[query_index] if query_index < len(all_metadatas) else []
            for index, distance in enumerate(distances):
                similarity = max(0.0, min(1.0, 1 - float(distance)))
                if similarity < threshold:
                    continue
                cleaned_document = TextCleaner.clean(documents[index])
                if not is_meaningful_chunk(cleaned_document):
                    continue
                meta = metadatas[index] or {}
                bucket.append({
                    "chunk_id": meta.get("chunk_id"),
                    "content": cleaned_document,
                    "similarity": round(similarity, 4),
                    "distance": round(float(distance), 4),
                    "retrieval_query": query,
                    **meta,
                })
                if len(bucket) >= per_query:
                    break
            if not bucket:
                bucket = self._lexical_fallback(query, filters, per_query)
            buckets.append(bucket)
        merged, seen_ids = [], set()
        for offset in range(max((len(bucket) for bucket in buckets), default=0)):
            for bucket in buckets:
                if offset >= len(bucket):
                    continue
                item = bucket[offset]
                chunk_key = (item.get("chunk_id"), item.get("retrieval_query"))
                if chunk_key in seen_ids:
                    continue
                seen_ids.add(chunk_key)
                merged.append(item)
                if len(merged) >= max_results:
                    return merged
        return merged

class FileParseService:
    @transaction.atomic
    def prepare_chunks(self, knowledge_file):
        config = {**get_config(), **knowledge_file.parse_config}
        knowledge_file.parse_status = "PARSING"; knowledge_file.parse_progress = 10; knowledge_file.error_message = ""; knowledge_file.save()
        pages = ParserFactory.create(knowledge_file.file_type).parse(knowledge_file.file.path)
        knowledge_file.parse_status = "CLEANING"; knowledge_file.parse_progress = 30; knowledge_file.save()
        records, index, char_count = [], 0, 0
        for page in pages:
            cleaned = TextCleaner.clean(page["text"]); char_count += len(cleaned)
            for content in split_text(cleaned, int(config.get("chunk_size", 800)), int(config.get("chunk_overlap", 120))):
                records.append(TextChunk(course=knowledge_file.course, knowledge_file=knowledge_file, chunk_index=index, content=content, content_hash=hashlib.sha256(content.encode()).hexdigest(), page_number=page["page"], char_count=len(content), vector_id=f"chunk-{knowledge_file.id}-{index}-{uuid.uuid4().hex[:8]}", metadata={"file_name": knowledge_file.original_name}))
                index += 1
        if not records: raise BusinessError("文件中没有可用于构建知识库的文本。", 40025)
        TextChunk.objects.filter(knowledge_file=knowledge_file).delete()
        TextChunk.objects.bulk_create(records, batch_size=200)
        knowledge_file.char_count = char_count; knowledge_file.chunk_count = len(records); knowledge_file.parse_status = "CHUNKING"; knowledge_file.parse_progress = 60; knowledge_file.save()
        return list(TextChunk.objects.filter(knowledge_file=knowledge_file))

    def run(self, task):
        file = task.knowledge_file
        vector_only = task.current_step == "等待重建向量" and file.chunks.exists()
        task.status = "RUNNING"; task.started_at = timezone.now(); task.heartbeat_at = timezone.now(); task.attempt_count += 1; task.save()
        try:
            if vector_only:
                # Chroma 索引可从 SQLite 中的原文本块恢复。不要重新分块，避免改变
                # TextChunk 主键并破坏章节、知识点等已经建立的关联。
                all_chunks = file.chunks.select_related("knowledge_file").order_by("chunk_index")
                total_chunks = all_chunks.count()
                # Worker 意外停止后从断点继续，不重复计算已成功写入的向量。
                chunks = list(all_chunks.exclude(vector_status="SUCCESS"))
                completed_before_resume = total_chunks - len(chunks)
                file.parse_status = "VECTORIZING"; file.parse_progress = 75; file.error_message = ""; file.save()
            else:
                if file.chunks.exists(): VectorService().delete_file(file.id)
                chunks = self.prepare_chunks(file)
                total_chunks = len(chunks)
                completed_before_resume = 0
            if task.cancel_requested: raise BusinessError("任务已取消", 40901, 409)
            file.parse_status = "VECTORIZING"; file.parse_progress = 75; file.save()
            task.progress = 75; task.current_step = "正在生成向量"; task.heartbeat_at = timezone.now(); task.save()
            def update_vector_progress(completed, total):
                task.progress = min(99, 75 + int(completed / total * 24))
                task.current_step = f"正在生成向量（{completed}/{total}）"
                task.heartbeat_at = timezone.now()
                task.save(update_fields=["progress", "current_step", "heartbeat_at"])

            def cancelled():
                return ParseTask.objects.filter(pk=task.pk, cancel_requested=True).exists()

            def update_resumable_progress(completed, _remaining_total):
                update_vector_progress(completed_before_resume + completed, total_chunks)

            def generation_waiting():
                from apps.questions.models import GenerationTask
                from apps.agents.models import AgentWorkflowRun
                return (
                    AgentWorkflowRun.objects.filter(status="WAITING", cancel_requested=False, priority__gte=70).exists()
                    or GenerationTask.objects.filter(status="WAITING", cancel_requested=False).exists()
                )

            VectorService().index(
                chunks,
                progress_callback=update_resumable_progress,
                cancel_callback=cancelled,
                yield_callback=generation_waiting,
            )
            file.parse_status = "SUCCESS"; file.parse_progress = 100; file.save()
            task.status = "SUCCESS"; task.progress = 100; task.current_step = "解析完成"; task.finished_at = timezone.now(); task.save()
            from apps.agents.services import create_knowledge_curation_workflow
            create_knowledge_curation_workflow(file)
        except TaskYield:
            # 保留已完成的向量，让Worker先处理等待中的出题任务。
            file.parse_status = "VECTORIZING"; file.error_message = ""; file.save(update_fields=["parse_status", "error_message", "updated_at"])
            task.status = "WAITING"; task.current_step = "等待重建向量"; task.error_message = ""; task.finished_at = None
            task.save(update_fields=["status", "current_step", "error_message", "finished_at"])
            logger.info("解析任务%s已保存断点，优先处理出题任务", task.id)
        except Exception as exc:
            status = "CANCELLED" if task.cancel_requested else "FAILED"
            file.parse_status = status if status == "FAILED" else "INTERRUPTED"; file.error_message = str(exc); file.save()
            task.status = status; task.error_message = str(exc); task.finished_at = timezone.now(); task.save()
            logger.exception("解析任务%s失败", task.id)
