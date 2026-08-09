import json
import logging
import os
import platform
import re
import uuid
from collections import Counter
from difflib import SequenceMatcher
from pathlib import Path
from django.conf import settings
from django.db import transaction
from common.exceptions import BusinessError
from apps.ai_service.services import OllamaService, get_prompt
from apps.questions.models import Question, GenerationTask
from apps.questions.services import (
    snapshot_question,
    find_similar,
    question_core_similarity,
    question_text_similarity,
    question_variation_template,
)
from .models import Paper, PaperSection, PaperQuestion, ExportRecord

logger = logging.getLogger("papers")

QUESTION_TYPE_ALIASES = {
    "single_choice": "single_choice", "单项选择题": "single_choice", "单选题": "single_choice",
    "multiple_choice": "multiple_choice", "多项选择题": "multiple_choice", "多选题": "multiple_choice",
    "judge": "judge", "判断题": "judge", "是非题": "judge",
    "fill_blank": "fill_blank", "填空题": "fill_blank",
    "term_explanation": "term_explanation", "名词解释题": "term_explanation", "名词解释": "term_explanation",
    "short_answer": "short_answer", "简答题": "short_answer",
    "essay": "essay", "论述题": "essay",
    "calculation": "calculation", "计算题": "calculation",
    "programming": "programming", "编程题": "programming", "程序设计题": "programming",
    "case_analysis": "case_analysis", "案例分析题": "case_analysis", "案例分析": "case_analysis",
}
QUESTION_TYPE_LABELS = {
    "single_choice": "选择题",
    "multiple_choice": "多项选择题",
    "judge": "判断题",
    "fill_blank": "填空题",
    "term_explanation": "名词解释题",
    "short_answer": "简答题",
    "essay": "论述题",
    "calculation": "计算题",
    "programming": "编程设计题",
    "case_analysis": "案例分析题",
}
DIFFICULTY_ALIASES = {
    "简单": "简单", "较易": "简单", "easy": "简单",
    "中等": "中等", "medium": "中等",
    "困难": "困难", "较难": "困难", "hard": "困难",
}


def _format_score(value):
    """导出时去掉无意义的小数，例如把2.00显示为2。"""
    return f"{float(value or 0):g}"


def _chinese_number(value):
    """把常见的大题序号转为中文，超出99时保留数字以避免错误转换。"""
    value = int(value)
    digits = "零一二三四五六七八九"
    if value < 10:
        return digits[value]
    if value < 20:
        return "十" + (digits[value % 10] if value % 10 else "")
    if value < 100:
        return digits[value // 10] + "十" + (digits[value % 10] if value % 10 else "")
    return str(value)


def section_display_title(section, section_index):
    """为新旧试卷生成统一的中文大题标题，同时保留用户自定义标题。"""
    raw_title = str(section.title or "").strip()
    title_without_order = re.sub(
        r"^(?:第\s*(?:\d+|[一二三四五六七八九十百]+)\s*部分|[一二三四五六七八九十百]+、)\s*",
        "",
        raw_title,
    ).strip()
    for code, label in QUESTION_TYPE_LABELS.items():
        title_without_order = re.sub(rf"\b{re.escape(code)}\b", label, title_without_order)

    if not title_without_order:
        question_types = {
            item.question_snapshot.get("question_type")
            for item in section.paper_questions.all()
            if item.question_snapshot.get("question_type")
        }
        title_without_order = QUESTION_TYPE_LABELS.get(next(iter(question_types)), "综合题") if len(question_types) == 1 else "综合题"
    return f"第{_chinese_number(section_index + 1)}部分 {title_without_order}"


def section_score_summary(section, questions=None):
    """分值只集中展示在大题标题；同分时同时说明每题分值。"""
    questions = list(questions if questions is not None else section.paper_questions.all())
    total_score = sum((item.score for item in questions), 0)
    summary = f"共{_format_score(total_score)}分"
    scores = {_format_score(item.score) for item in questions}
    if questions and len(scores) == 1:
        summary += f"，每题{next(iter(scores))}分"
    return summary


def _docx_font_name():
    system = platform.system()
    if system == "Darwin":
        # macOS系统自带完整中英文字符集，Word与LibreOffice均可识别。
        return "Arial Unicode MS"
    if system == "Windows":
        return "宋体"
    return "Noto Sans CJK SC"


def _number(value, field_name, minimum=0):
    try:
        result = float(value)
    except (TypeError, ValueError) as exc:
        raise BusinessError(f"{field_name}必须是数字。", 40053) from exc
    if result < minimum:
        raise BusinessError(f"{field_name}不能小于{minimum}。", 40053)
    return result


def _boolean(value, default=False):
    if value is None:
        return default
    if isinstance(value, str):
        return value.strip().lower() in {"1", "true", "yes", "on", "是", "允许"}
    return bool(value)


def _normalize_type_config(raw_config):
    """兼容本地模型常见的中英文字段和对象/列表两种写法。"""
    if isinstance(raw_config, dict):
        raw_config = [
            ({"type": key, **value} if isinstance(value, dict) else {"type": key, "count": value})
            for key, value in raw_config.items()
        ]
    if not isinstance(raw_config, list) or not raw_config:
        raise BusinessError("AI组卷规则中没有有效的题型配置。", 40054)
    normalized = []
    for index, item in enumerate(raw_config, 1):
        if not isinstance(item, dict):
            raise BusinessError(f"第{index}项题型配置格式不正确。", 40054)
        raw_type = item.get("type") or item.get("question_type") or item.get("题型")
        question_type = QUESTION_TYPE_ALIASES.get(str(raw_type).strip())
        if not question_type:
            raise BusinessError(f"AI返回了不支持的题型“{raw_type}”，请修改要求后重试。", 40055)
        count = int(_number(item.get("count", item.get("quantity", item.get("数量"))), f"第{index}项题型数量", 1))
        score = _number(item.get("score_each", item.get("score", item.get("每题分值", 1))), f"第{index}项每题分值", 0.5)
        normalized.append({
            "type": question_type,
            "count": count,
            "score_each": int(score) if score.is_integer() else score,
            **({"title": item["title"]} if item.get("title") else {}),
        })
    return normalized


def _normalize_difficulty_ratio(raw_ratios):
    if not raw_ratios:
        return {"简单": 0.3, "中等": 0.5, "困难": 0.2}
    if isinstance(raw_ratios, list):
        if all(isinstance(item, (int, float)) for item in raw_ratios) and len(raw_ratios) == 3:
            raw_ratios = dict(zip(["简单", "中等", "困难"], raw_ratios))
        else:
            converted = {}
            for item in raw_ratios:
                if not isinstance(item, dict):
                    raise BusinessError("难度比例列表格式不正确。", 40056)
                name = item.get("difficulty") or item.get("name") or item.get("level") or item.get("难度")
                value = item.get("ratio", item.get("value", item.get("percentage", item.get("比例"))))
                if name is None and len(item) == 1:
                    name, value = next(iter(item.items()))
                converted[name] = value
            raw_ratios = converted
    if not isinstance(raw_ratios, dict):
        raise BusinessError("难度比例必须是对象或列表。", 40056)
    normalized = {"简单": 0.0, "中等": 0.0, "困难": 0.0}
    for raw_name, raw_value in raw_ratios.items():
        name = DIFFICULTY_ALIASES.get(str(raw_name).strip().lower())
        if not name:
            raise BusinessError(f"不支持的难度名称“{raw_name}”。", 40056)
        normalized[name] += _number(raw_value, f"{raw_name}难度比例")
    total = sum(normalized.values())
    if total > 1.5:  # 兼容30/50/20这种百分数写法
        normalized = {key: value / 100 for key, value in normalized.items()}
        total = sum(normalized.values())
    if total <= 0 or abs(total - 1.0) > 0.02:
        raise BusinessError(f"难度比例之和必须等于100%，当前为{total * 100:g}%。", 40051)
    # 模型可能输出0.33/0.33/0.33，在误差范围内归一化为精确的1。
    return {key: round(value / total, 6) for key, value in normalized.items()}


def normalize_paper_rule(data):
    if not isinstance(data, dict):
        raise BusinessError("AI返回的组卷规则不是JSON对象。", 40057)
    nested = data.get("rule") or data.get("paper_rule")
    if isinstance(nested, dict):
        data = {**nested, **{key: value for key, value in data.items() if key not in {"rule", "paper_rule"}}}
    result = dict(data)
    result["name"] = str(data.get("name") or data.get("paper_name") or data.get("试卷名称") or "AI辅助组卷").strip()
    result["paper_type"] = data.get("paper_type") or "AI辅助组卷"
    if data.get("duration") is not None or data.get("exam_time") is not None:
        result["duration"] = int(_number(data.get("duration", data.get("exam_time")), "考试时间", 1))
    else:
        result.pop("duration", None)
        result.pop("exam_time", None)
    result["type_config"] = _normalize_type_config(data.get("type_config") or data.get("question_types"))
    calculated_score = sum(item["count"] * item["score_each"] for item in result["type_config"])
    target_score = _number(data.get("target_score", data.get("total_score", calculated_score)), "目标总分", 1)
    result["target_score"] = int(target_score) if target_score.is_integer() else target_score
    result["difficulty_ratio"] = _normalize_difficulty_ratio(data.get("difficulty_ratio") or data.get("difficulty_ratios"))
    for key, default in (("allow_similar", False), ("prefer_unused", True), ("allow_ai_fill", False)):
        result[key] = _boolean(data.get(key), default)
    return result


def _resolve_scope_ids(values, model, course_id):
    """把模型返回的章节/知识点名称转成当前课程的真实ID。"""
    if not isinstance(values, list):
        return []
    resolved = []
    for value in values:
        if isinstance(value, int) or (isinstance(value, str) and value.strip().isdigit()):
            item_id = int(value)
            if model.objects.filter(pk=item_id, course_id=course_id).exists():
                resolved.append(item_id)
            continue
        name = str(value or "").strip()
        if name:
            item_id = model.objects.filter(course_id=course_id, name__iexact=name).values_list("id", flat=True).first()
            if item_id:
                resolved.append(item_id)
    return list(dict.fromkeys(resolved))

def recalculate_paper(paper):
    total = sum(item.score for item in paper.paper_questions.all())
    paper.total_score = total; paper.save(update_fields=["total_score", "updated_at"])
    for section in paper.sections.all():
        section.score = sum(x.score for x in section.paper_questions.all()); section.save(update_fields=["score"])
    return total


def paper_question_similarity(first, second):
    """判断两道题是否过于接近；该规则只用于同一张试卷。"""
    text_score = question_text_similarity(first.stem, second.stem)
    first_template = question_variation_template(first.stem)
    second_template = question_variation_template(second.stem)
    template_score = SequenceMatcher(None, first_template, second_template).ratio() if first_template and second_template else 0.0
    core_score, core_containment = question_core_similarity(first.stem, second.stem)
    meaningful_template = min(len(first_template), len(second_template)) >= 6
    similar = (
        text_score >= 0.88
        or (meaningful_template and template_score >= 0.92)
        or core_score >= 0.86
        or core_containment >= 0.78
    )
    return similar, max(text_score, template_score, core_score, core_containment)


def find_similar_question_in_paper(question, existing_questions):
    for existing in existing_questions:
        similar, score = paper_question_similarity(question, existing)
        if similar:
            return existing, score
    return None, 0.0


def coverage_aware_selection(candidates, count, allow_similar=False, existing_questions=None):
    """贪心选题：优先补足新知识点/章节，并确保整张试卷没有相近题。"""
    selected, knowledge_counts, chapter_counts = [], Counter(), Counter()
    existing_questions = list(existing_questions or [])
    remaining = list(candidates)
    while remaining and len(selected) < count:
        def score(question):
            knowledge_gain = 3.0 if question.knowledge_point_id and not knowledge_counts[question.knowledge_point_id] else 0.0
            chapter_gain = 1.5 if question.chapter_id and not chapter_counts[question.chapter_id] else 0.0
            unused_bonus = 2.0 / (1 + question.use_count)
            concentration_penalty = knowledge_counts[question.knowledge_point_id] * 0.8 if question.knowledge_point_id else 0.0
            similarity_penalty = 0.0
            comparison_pool = existing_questions + selected
            if comparison_pool:
                similarity_penalty = max(paper_question_similarity(question, item)[1] for item in comparison_pool) * 4
            return knowledge_gain + chapter_gain + unused_bonus - concentration_penalty - similarity_penalty
        best = max(remaining, key=score)
        if find_similar_question_in_paper(best, existing_questions + selected)[0] is not None:
            remaining.remove(best)
            continue
        selected.append(best); remaining.remove(best)
        knowledge_counts[best.knowledge_point_id] += 1
        chapter_counts[best.chapter_id] += 1
    return selected

@transaction.atomic
def rule_generate(data):
    data = normalize_paper_rule(data)
    ratios = data["difficulty_ratio"]
    course_id = data.get("course") or data.get("course_id")
    if not course_id:
        raise BusinessError("请选择组卷课程。", 40059)
    from apps.knowledge.models import Chapter, KnowledgePoint
    data["chapter_ids"] = _resolve_scope_ids(data.get("chapter_ids"), Chapter, course_id)
    data["knowledge_point_ids"] = _resolve_scope_ids(data.get("knowledge_point_ids"), KnowledgePoint, course_id)
    paper_type = data.get("paper_type", "规则组卷")
    # AI辅助组卷不要求考试时间，0表示预览和导出时不显示。
    duration = data.get("duration", 0 if paper_type == "AI辅助组卷" else 90)
    paper = Paper.objects.create(course_id=course_id, name=data["name"], paper_type=paper_type, duration=duration, target_score=data.get("target_score", 100), config=data)
    shortages = []
    paper_selected = []
    for section_index, rule in enumerate(data.get("type_config", [])):
        qtype, count, score_each = rule["type"], int(rule["count"]), rule["score_each"]
        qs = Question.objects.filter(course_id=course_id, question_type=qtype, review_status="APPROVED", is_deleted=False)
        if data.get("chapter_ids"): qs = qs.filter(chapter_id__in=data["chapter_ids"])
        if data.get("knowledge_point_ids"): qs = qs.filter(knowledge_point_id__in=data["knowledge_point_ids"])
        if data.get("prefer_unused", True): qs = qs.order_by("use_count", "id")
        else: qs = qs.order_by("id")
        candidates = list(qs[:max(count * 10, 100)])
        selected = []
        difficulty_groups = {"简单": ["简单", "较易"], "中等": ["中等"], "困难": ["较难", "困难"]}
        if ratios:
            allocations, remaining = {}, count
            ordered_ratios = list(ratios.items())
            for idx, (name, ratio) in enumerate(ordered_ratios):
                allocation = remaining if idx == len(ordered_ratios) - 1 else int(count * float(ratio))
                allocations[name] = max(0, allocation); remaining -= allocation
            for name, allocation in allocations.items():
                selected.extend([q for q in candidates if q.difficulty in difficulty_groups.get(name, [name]) and q not in selected][:allocation])
        if len(selected) < count: selected.extend([q for q in candidates if q not in selected][:count - len(selected)])
        preferred = selected + [q for q in candidates if q not in selected]
        selected = coverage_aware_selection(preferred, count, existing_questions=paper_selected)
        paper_selected.extend(selected)
        if len(selected) < count: shortages.append({"type": qtype, "required": count, "available": len(selected), "missing": count - len(selected)})
        section = PaperSection.objects.create(
            paper=paper,
            title=f"第{_chinese_number(section_index + 1)}部分 {QUESTION_TYPE_LABELS[qtype]}",
            sort_order=section_index,
        )
        PaperQuestion.objects.bulk_create([PaperQuestion(paper=paper, section=section, question=q, sort_order=i, score=score_each, question_snapshot=snapshot_question(q)) for i, q in enumerate(selected)])
        Question.objects.filter(id__in=[q.id for q in selected]).update(use_count=__import__("django.db.models", fromlist=["F"]).F("use_count") + 1)
    recalculate_paper(paper)
    if shortages and data.get("allow_ai_fill"):
        task_ids = []
        from apps.agents.services import create_question_workflow
        for shortage in shortages:
            task = GenerationTask.objects.create(course_id=course_id, status="WAITING", total_count=shortage["missing"], config={"course":course_id, "count":shortage["missing"], "question_types":[shortage["type"]], "difficulty":"中等", "score":next((x["score_each"] for x in data.get("type_config", []) if x["type"] == shortage["type"]), 1), "scenario":"试卷补题", "strict":True})
            create_question_workflow(task, "STANDARD")
            task_ids.append(task.id)
        paper.config = {**paper.config, "ai_fill_task_ids": task_ids}; paper.save(update_fields=["config"])
    from apps.agents.services import record_paper_selection_workflow
    record_paper_selection_workflow(paper, shortages)
    return paper, shortages

def parse_natural_rule(text, course_id=None):
    if not str(text).strip():
        raise BusinessError("请输入组卷要求。", 40058)
    prompt = get_prompt("paper_rule", "把自然语言试卷要求转换成结构化JSON，包含name、target_score、type_config、difficulty_ratio，不解析考试时间。")
    result = OllamaService().chat_json([{"role": "system", "content": prompt}, {"role": "user", "content": text}], purpose="paper_rule")
    result = normalize_paper_rule(result)
    # 即使模型没有遵守提示词，也不把考试时间带入AI组卷规则。
    result.pop("duration", None)
    result.pop("exam_time", None)
    result["paper_type"] = "AI辅助组卷"
    if course_id:
        from apps.knowledge.models import Chapter, KnowledgePoint
        result["course_id"] = course_id
        result["chapter_ids"] = _resolve_scope_ids(result.get("chapter_ids"), Chapter, course_id)
        result["knowledge_point_ids"] = _resolve_scope_ids(result.get("knowledge_point_ids"), KnowledgePoint, course_id)
    return result

def quality_analysis(paper):
    items = list(paper.paper_questions.all())
    type_counts = Counter(x.question_snapshot.get("question_type") for x in items)
    difficulty_counts = Counter(x.question_snapshot.get("difficulty") for x in items)
    knowledge = [x.question_snapshot.get("knowledge_point") for x in items if x.question_snapshot.get("knowledge_point")]
    issues, suggestions = [], []
    total = float(sum(x.score for x in items))
    if abs(total - float(paper.target_score)) > 0.01: issues.append(f"当前总分为{total:g}分，与目标总分{float(paper.target_score):g}分不一致")
    no_answer = sum(not x.question_snapshot.get("answer") for x in items)
    no_analysis = sum(not x.question_snapshot.get("analysis") for x in items)
    pending = sum((x.question.review_status if x.question else "") != "APPROVED" for x in items)
    if no_answer: issues.append(f"有{no_answer}道题缺少答案")
    if no_analysis: issues.append(f"有{no_analysis}道题缺少详细解析")
    if pending: issues.append(f"有{pending}道题不是已审核状态")
    if knowledge and Counter(knowledge).most_common(1)[0][1] / len(items) > 0.4: suggestions.append("知识点分布较集中，建议补充其他知识点题目")
    completeness = max(0, 100 - no_answer * 10 - no_analysis * 5 - pending * 10)
    score = round((90 + min(100, len(set(knowledge)) * 10) + min(100, len(type_counts) * 20) + 95 + completeness) / 5, 1)
    result = {"total_score": score, "grade": "优秀" if score >= 90 else "良好" if score >= 80 else "合格" if score >= 60 else "需改进", "dimensions": {"difficulty_balance": 90 if len(difficulty_counts) >= 3 else 70, "knowledge_coverage": min(100, len(set(knowledge)) * 10), "question_type_balance": min(100, len(type_counts) * 20), "duplication_control": 95, "completeness": completeness}, "statistics": {"question_types": type_counts, "difficulties": difficulty_counts, "total_questions": len(items), "current_score": total}, "issues": issues, "suggestions": suggestions}
    paper.quality_score = score; paper.save(update_fields=["quality_score"]); return result

def _add_docx_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run(); begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, end])

class ExportService:
    def export(self, paper, export_type, file_format):
        record = ExportRecord.objects.create(paper=paper, export_type=export_type, file_format=file_format, status="RUNNING")
        export_dir = settings.MEDIA_ROOT / "exports"; export_dir.mkdir(parents=True, exist_ok=True)
        filename = f"{paper.id}_{export_type}_{uuid.uuid4().hex[:8]}.{file_format}"; path = export_dir / filename
        try:
            if file_format == "docx": self._docx(paper, export_type, path)
            elif file_format == "pdf": self._pdf(paper, export_type, path)
            else: raise BusinessError("仅支持docx和pdf导出。", 40052)
            record.file_name = filename; record.file_path = str(path.relative_to(settings.MEDIA_ROOT)); record.file_size = path.stat().st_size; record.status = "SUCCESS"; record.save()
            paper.status = "EXPORTED"; paper.save(update_fields=["status"]); return record
        except Exception as exc:
            record.status = "FAILED"; record.error_message = str(exc); record.save(); logger.exception("试卷导出失败"); raise BusinessError(f"导出失败：{exc}", 50051, 500) from exc
    def _docx(self, paper, export_type, path):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.shared import Cm, Pt
        doc = Document(); section = doc.sections[0]; section.page_height = Cm(29.7); section.page_width = Cm(21); section.top_margin = section.bottom_margin = Cm(2); section.left_margin = section.right_margin = Cm(2.2)
        styles = doc.styles
        font_name = _docx_font_name()
        for style_name in ["Normal", "Title", "Heading 1", "Heading 2"]:
            style = styles[style_name]
            style.font.name = font_name
            for font_scope in ["w:ascii", "w:hAnsi", "w:eastAsia"]:
                style._element.rPr.rFonts.set(qn(font_scope), font_name)
        styles["Normal"].font.size = Pt(11)
        header = section.header.paragraphs[0]; header.text = paper.school_name or "知识库智能出题与组卷系统"; header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER; _add_docx_page_number(footer)
        title = doc.add_heading(paper.name, 0); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info = doc.add_table(rows=2, cols=4); info.style = "Table Grid"
        values = ["专业", paper.major, "班级", paper.class_name, "姓名", "____________", "学号", "____________"]
        for i, cell in enumerate([c for row in info.rows for c in row.cells]): cell.text = str(values[i] or "")
        info_parts = []
        if paper.duration:
            info_parts.append(f"考试时间：{paper.duration}分钟")
        info_parts.append(f"总分：{_format_score(paper.total_score)}分")
        doc.add_paragraph("    ".join(info_parts))
        if paper.instructions: doc.add_paragraph(f"考试说明：{paper.instructions}")
        number = 1
        for section_index, sec in enumerate(paper.sections.prefetch_related("paper_questions")):
            section_questions = list(sec.paper_questions.all())
            doc.add_heading(f"{section_display_title(sec, section_index)}（{section_score_summary(sec, section_questions)}）", level=1)
            if sec.description: doc.add_paragraph(sec.description)
            for item in section_questions:
                q = item.question_snapshot; doc.add_paragraph(f"{number}. {q.get('stem', '')}")
                for opt in q.get("options", []): doc.add_paragraph(f"   {opt.get('label')}. {opt.get('content')}")
                if export_type != "student":
                    doc.add_paragraph("答案：" + "、".join(map(str, q.get("answer", []))))
                    if q.get("scoring_points"): doc.add_paragraph("评分要点：" + "；".join(q["scoring_points"]))
                if export_type == "analysis":
                    doc.add_paragraph("解析：" + (q.get("analysis") or "暂无")); doc.add_paragraph(f"知识点：{q.get('knowledge_point') or '未关联'}  难度：{q.get('difficulty')}  来源：{q.get('source_summary') or '未记录'}")
                elif export_type == "student" and q.get("question_type") not in ["single_choice", "multiple_choice", "judge"]:
                    doc.add_paragraph("\n" * 2)
                number += 1
        doc.save(path)
    def _pdf(self, paper, export_type, path):
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        candidates = ["/System/Library/Fonts/STHeiti Light.ttc", "/System/Library/Fonts/PingFang.ttc", "C:/Windows/Fonts/simsun.ttc", "C:/Windows/Fonts/msyh.ttc"]
        font_path = next((p for p in candidates if os.path.exists(p)), None)
        if not font_path: raise BusinessError("未找到可用中文字体，请在系统中安装宋体或微软雅黑后重试PDF导出。", 50052, 500)
        pdfmetrics.registerFont(TTFont("Chinese", font_path, subfontIndex=0))
        styles = getSampleStyleSheet(); body = ParagraphStyle("ChineseBody", parent=styles["BodyText"], fontName="Chinese", fontSize=10.5, leading=17); title = ParagraphStyle("ChineseTitle", parent=body, fontSize=18, leading=26, alignment=TA_CENTER); heading = ParagraphStyle("ChineseHeading", parent=body, fontSize=13, leading=21, spaceBefore=10)
        doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=2*cm, rightMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        score_line = f"总分：{_format_score(paper.total_score)}分"
        if paper.duration:
            score_line = f"考试时间：{paper.duration}分钟　{score_line}"
        story = [Paragraph(paper.name, title), Spacer(1, 8), Table([[f"专业：{paper.major}", f"班级：{paper.class_name}"], ["姓名：____________", "学号：____________"]], colWidths=[8.5*cm, 8.5*cm], style=TableStyle([("FONTNAME", (0,0), (-1,-1), "Chinese"), ("GRID", (0,0), (-1,-1), .5, colors.grey), ("PADDING", (0,0), (-1,-1), 6)])), Spacer(1, 8), Paragraph(score_line, body)]
        if paper.instructions: story.append(Paragraph("考试说明：" + paper.instructions, body))
        number = 1
        for section_index, sec in enumerate(paper.sections.prefetch_related("paper_questions")):
            section_questions = list(sec.paper_questions.all())
            story.append(Paragraph(f"{section_display_title(sec, section_index)}（{section_score_summary(sec, section_questions)}）", heading))
            for item in section_questions:
                q = item.question_snapshot; story.append(Paragraph(f"{number}. {q.get('stem','')}", body))
                for opt in q.get("options", []): story.append(Paragraph(f"　{opt.get('label')}. {opt.get('content')}", body))
                if export_type != "student": story.append(Paragraph("答案：" + "、".join(map(str, q.get("answer", []))), body))
                if export_type == "analysis": story.extend([Paragraph("解析：" + (q.get("analysis") or "暂无"), body), Paragraph("知识点：" + (q.get("knowledge_point") or "未关联") + "　来源：" + (q.get("source_summary") or "未记录"), body)])
                elif export_type == "student" and q.get("question_type") not in ["single_choice", "multiple_choice", "judge"]: story.append(Spacer(1, 2*cm))
                number += 1
        doc.build(story, onFirstPage=self._page_number, onLaterPages=self._page_number)
    @staticmethod
    def _page_number(canvas, doc):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import cm
        canvas.saveState(); canvas.setFont("Chinese", 9); canvas.drawCentredString(A4[0]/2, 1*cm, f"第 {doc.page} 页"); canvas.restoreState()
